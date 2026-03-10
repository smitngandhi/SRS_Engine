from __future__ import annotations

"""
parse_service.py
────────────────
Parser Agent — accuracy-first, handles mixed document types.

Extraction chains:
  PDF  → pdfplumber (primary)
           └─ pdfminer.six (fallback on extraction failure)
               └─ pytesseract + pdf2image (OCR fallback for scanned PDFs)
       → camelot  (parallel, for tables only)

  DOCX → python-docx (primary, preserves heading styles + tables)
           └─ mammoth → BeautifulSoup (fallback for malformed DOCX)

Output: UnifiedDocumentJSON saved to parsed_docs/{user_id}/{file_id}.json
"""

import json
import re
from datetime import datetime, timezone
from pathlib import Path

from fastapi import HTTPException

from srs_engine.schemas.upload_schemas.parse_schema import (
    ParseMetadata,
    ParseResponse,
    ParsedSection,
    ParsedTable,
    UnifiedDocumentJSON,
)

# ── Constants ────────────────────────────────────────────────────────────────

PARSED_DOCS_ROOT = Path("./parsed_docs")

# Regex patterns for section heading detection in plain text / PDF
_NUMBERED_HEADING = re.compile(
    r"^(\d+(?:\.\d+){0,3})\s{1,4}([A-Z][^\n]{2,80})$", re.MULTILINE
)
_CAPS_HEADING = re.compile(r"^([A-Z][A-Z\s\-]{4,60})$", re.MULTILINE)


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _parsed_path(user_id: str, file_id: str) -> Path:
    directory = PARSED_DOCS_ROOT / user_id
    directory.mkdir(parents=True, exist_ok=True)
    return directory / f"{file_id}.json"


# ─────────────────────────────────────────────────────────────────────────────
# PDF PARSING
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf_text_pdfplumber(file_path: Path) -> tuple[str, int, list[ParsedTable]]:
    """Primary PDF extractor. Returns (text, page_count, tables)."""
    import pdfplumber

    full_text_parts: list[str] = []
    all_tables: list[ParsedTable] = []
    page_count = 0

    with pdfplumber.open(str(file_path)) as pdf:
        page_count = len(pdf.pages)
        for page_idx, page in enumerate(pdf.pages):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            full_text_parts.append(text)

            # Extract tables via pdfplumber
            raw_tables = page.extract_tables() or []
            for t_idx, raw_table in enumerate(raw_tables):
                if not raw_table:
                    continue
                headers = [str(c or "").strip() for c in raw_table[0]]
                rows = [
                    [str(cell or "").strip() for cell in row]
                    for row in raw_table[1:]
                ]
                all_tables.append(ParsedTable(
                    table_index=len(all_tables),
                    headers=headers,
                    rows=rows,
                    extraction_method="pdfplumber",
                ))

    return "\n".join(full_text_parts), page_count, all_tables


def _extract_pdf_tables_camelot(file_path: Path) -> list[ParsedTable]:
    """Parallel table extractor using camelot. Merges/replaces pdfplumber tables."""
    try:
        import camelot
        tables_lattice = camelot.read_pdf(str(file_path), flavor="lattice", pages="all")
        tables_stream  = camelot.read_pdf(str(file_path), flavor="stream",  pages="all")

        result: list[ParsedTable] = []
        for t_idx, table in enumerate(list(tables_lattice) + list(tables_stream)):
            df = table.df
            if df.empty:
                continue
            headers = [str(c).strip() for c in df.iloc[0].tolist()]
            rows    = [[str(cell).strip() for cell in row] for row in df.iloc[1:].values.tolist()]
            result.append(ParsedTable(
                table_index=t_idx,
                headers=headers,
                rows=rows,
                extraction_method="camelot",
            ))
        return result
    except Exception:
        return []  # camelot is optional — never block the pipeline


def _extract_pdf_text_pdfminer(file_path: Path) -> tuple[str, int]:
    """Fallback PDF extractor using pdfminer.six."""
    from pdfminer.high_level import extract_text, extract_pages
    from pdfminer.layout import LTPage

    text = extract_text(str(file_path)) or ""
    page_count = sum(1 for _ in extract_pages(str(file_path)) if isinstance(_, LTPage))
    return text, page_count


def _extract_pdf_text_ocr(file_path: Path) -> tuple[str, int]:
    """Last-resort OCR for scanned PDFs. Uses pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(str(file_path), dpi=300)
        pages_text = [pytesseract.image_to_string(img, lang="eng") for img in images]
        return "\n".join(pages_text), len(images)
    except Exception as exc:
        raise RuntimeError(f"OCR extraction failed: {exc}") from exc


def _is_text_empty(text: str) -> bool:
    """True if the extracted text is essentially blank (scanned PDF)."""
    return len(text.strip()) < 100


def parse_pdf(file_path: Path) -> dict:
    """
    Full PDF parse chain.
    Returns dict with keys: text, page_count, tables, extractor, fallback_used,
                             fallback_extractor, ocr_used, warnings
    """
    warnings: list[str] = []
    fallback_used = False
    fallback_extractor = None
    ocr_used = False

    # ── Primary: pdfplumber ──────────────────────────────
    try:
        text, page_count, plumber_tables = _extract_pdf_text_pdfplumber(file_path)
        extractor = "pdfplumber"
    except Exception as e:
        warnings.append(f"pdfplumber failed: {e}")
        text, page_count, plumber_tables = "", 0, []
        extractor = "pdfplumber"

    # ── Fallback: pdfminer.six ───────────────────────────
    if _is_text_empty(text):
        try:
            text, page_count = _extract_pdf_text_pdfminer(file_path)
            fallback_used = True
            fallback_extractor = "pdfminer.six"
            warnings.append("pdfplumber returned empty text; fell back to pdfminer.six")
        except Exception as e:
            warnings.append(f"pdfminer.six failed: {e}")

    # ── Last resort: OCR ─────────────────────────────────
    if _is_text_empty(text):
        try:
            text, page_count = _extract_pdf_text_ocr(file_path)
            ocr_used = True
            fallback_used = True
            fallback_extractor = "pytesseract"
            warnings.append("Text-based extraction empty; used OCR (pytesseract)")
        except Exception as e:
            warnings.append(f"OCR failed: {e}")

    if _is_text_empty(text):
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from this PDF. It may be corrupted or fully image-based without readable layers."
        )

    # ── Parallel: camelot for tables ─────────────────────
    camelot_tables = _extract_pdf_tables_camelot(file_path)
    # Prefer camelot tables if found (higher accuracy), else keep pdfplumber's
    tables = camelot_tables if camelot_tables else plumber_tables
    if camelot_tables:
        warnings.append(f"camelot found {len(camelot_tables)} tables (used over pdfplumber tables)")

    return dict(
        text=text,
        page_count=page_count,
        tables=tables,
        extractor=extractor,
        fallback_used=fallback_used,
        fallback_extractor=fallback_extractor,
        ocr_used=ocr_used,
        warnings=warnings,
    )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX PARSING
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx_python_docx(file_path: Path) -> tuple[str, list[ParsedSection], list[ParsedTable]]:
    """
    Primary DOCX extractor.
    Walks paragraphs preserving Heading 1/2/3 styles to build section tree.
    """
    import docx

    doc = docx.Document(str(file_path))
    sections: list[ParsedSection] = []
    tables: list[ParsedTable] = []
    raw_lines: list[str] = []

    # Section stack for building hierarchy
    stack: list[ParsedSection] = []   # [level-1-section, level-2-section, ...]

    def current_parent(level: int) -> list[ParsedSection]:
        """Return the subsections list of the appropriate parent."""
        for i in range(level - 2, -1, -1):
            if i < len(stack):
                return stack[i].subsections
        return sections

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        raw_lines.append(text)

        # Detect heading level from style
        level = None
        if "Heading 1" in style_name:
            level = 1
        elif "Heading 2" in style_name:
            level = 2
        elif "Heading 3" in style_name:
            level = 3
        elif "Heading 4" in style_name:
            level = 4

        if level is not None:
            # Extract numeric prefix if present (e.g. "1.2 Purpose")
            m = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)", text)
            section_id = m.group(1) if m else f"h{level}-{len(sections)}"
            heading = m.group(2) if m else text

            new_section = ParsedSection(
                section_id=section_id,
                heading=heading,
                level=level,
                content="",
            )

            # Trim stack to current level
            stack = stack[:level - 1]
            current_parent(level).append(new_section)
            stack.append(new_section)

        else:
            # Body text — append to the innermost open section
            if stack:
                stack[-1].content += ("\n" if stack[-1].content else "") + text
            # else it's preamble text before any heading — attach to raw_text only

    # ── Extract tables ────────────────────────────────────
    for t_idx, table in enumerate(doc.tables):
        rows_data: list[list[str]] = []
        for row in table.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if not rows_data:
            continue
        headers = rows_data[0]
        data_rows = rows_data[1:]
        tables.append(ParsedTable(
            table_index=t_idx,
            headers=headers,
            rows=data_rows,
            extraction_method="python-docx",
        ))

    raw_text = "\n".join(raw_lines)
    return raw_text, sections, tables


def _extract_docx_mammoth(file_path: Path) -> tuple[str, list[ParsedSection], list[ParsedTable]]:
    """
    Fallback DOCX extractor via mammoth → HTML → BeautifulSoup.
    Handles malformed DOCX or non-standard heading styles.
    """
    import mammoth
    from bs4 import BeautifulSoup

    with open(str(file_path), "rb") as f:
        result = mammoth.convert_to_html(f)

    html = result.value
    soup = BeautifulSoup(html, "html.parser")

    sections: list[ParsedSection] = []
    tables: list[ParsedTable] = []
    raw_lines: list[str] = []
    stack: list[ParsedSection] = []

    def current_parent(level: int) -> list[ParsedSection]:
        for i in range(level - 2, -1, -1):
            if i < len(stack):
                return stack[i].subsections
        return sections

    for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "table"]):
        tag = el.name

        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            text = el.get_text(strip=True)
            if not text:
                continue
            raw_lines.append(text)
            m = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)", text)
            section_id = m.group(1) if m else f"h{level}-{len(sections)}"
            heading = m.group(2) if m else text
            new_section = ParsedSection(
                section_id=section_id,
                heading=heading,
                level=level,
                content="",
            )
            stack = stack[:level - 1]
            current_parent(level).append(new_section)
            stack.append(new_section)

        elif tag == "p":
            text = el.get_text(strip=True)
            if not text:
                continue
            raw_lines.append(text)
            if stack:
                stack[-1].content += ("\n" if stack[-1].content else "") + text

        elif tag == "table":
            t_rows = []
            for tr in el.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    t_rows.append(cells)
            if t_rows:
                tables.append(ParsedTable(
                    table_index=len(tables),
                    headers=t_rows[0],
                    rows=t_rows[1:],
                    extraction_method="mammoth",
                ))

    return "\n".join(raw_lines), sections, tables


def parse_docx(file_path: Path) -> dict:
    """
    Full DOCX parse chain.
    Returns dict with keys: text, sections, tables, extractor,
                             fallback_used, fallback_extractor, warnings
    """
    warnings: list[str] = []
    fallback_used = False
    fallback_extractor = None

    try:
        text, sections, tables = _extract_docx_python_docx(file_path)
        extractor = "python-docx"

        # If no sections detected, structure is non-standard — try mammoth
        if not sections:
            warnings.append("python-docx found no heading-based sections; trying mammoth")
            raise ValueError("no sections")

    except Exception as e:
        warnings.append(f"python-docx primary failed ({e}); falling back to mammoth")
        try:
            text, sections, tables = _extract_docx_mammoth(file_path)
            extractor = "mammoth"
            fallback_used = True
            fallback_extractor = "mammoth"
        except Exception as e2:
            raise HTTPException(
                status_code=422,
                detail=f"Could not parse DOCX: {e2}"
            ) from e2

    return dict(
        text=text,
        sections=sections,
        tables=tables,
        extractor=extractor,
        fallback_used=fallback_used,
        fallback_extractor=fallback_extractor,
        ocr_used=False,
        warnings=warnings,
    )


# ─────────────────────────────────────────────────────────────────────────────
# SECTION DETECTION FROM PLAIN TEXT (for flat PDF output)
# ─────────────────────────────────────────────────────────────────────────────

def _detect_sections_from_text(text: str) -> list[ParsedSection]:
    """
    When the PDF extractor returns raw text without layout info,
    use regex heuristics to detect numbered sections.
    """
    sections: list[ParsedSection] = []
    lines = text.splitlines()
    current_section: ParsedSection | None = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        m = re.match(r"^(\d+(?:\.\d+){0,3})\s{1,4}(.{3,80})$", line)
        if m:
            section_id = m.group(1)
            heading    = m.group(2).strip()
            level      = section_id.count(".") + 1

            new_sec = ParsedSection(
                section_id=section_id,
                heading=heading,
                level=level,
                content="",
            )

            if level == 1:
                sections.append(new_sec)
                current_section = new_sec
            else:
                # Attach to parent
                parts = section_id.split(".")
                parent_id = ".".join(parts[:-1])
                parent = _find_section(sections, parent_id)
                if parent:
                    parent.subsections.append(new_sec)
                else:
                    sections.append(new_sec)
                current_section = new_sec
        else:
            if current_section is not None:
                current_section.content += ("\n" if current_section.content else "") + line

    return sections


def _find_section(sections: list[ParsedSection], section_id: str) -> ParsedSection | None:
    for sec in sections:
        if sec.section_id == section_id:
            return sec
        found = _find_section(sec.subsections, section_id)
        if found:
            return found
    return None


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC SERVICE FUNCTION
# ─────────────────────────────────────────────────────────────────────────────

async def parse_uploaded_file(user_id: str, file_id: str, storage_path: str, original_filename: str, file_type: str) -> ParseResponse:
    """
    Entry point called by the router.
    Reads the file from disk, runs the appropriate parse chain,
    builds UnifiedDocumentJSON, saves it, and returns ParseResponse.
    """
    file_path = Path(storage_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Uploaded file not found on disk.")

    # ── Run the correct parse chain ───────────────────────
    if file_type == "pdf":
        result = parse_pdf(file_path)

        # Build section tree from raw text if pdfplumber didn't give us structure
        sections = _detect_sections_from_text(result["text"])
        tables   = result["tables"]

    elif file_type == "docx":
        result   = parse_docx(file_path)
        sections = result["sections"]
        tables   = result["tables"]

        # If sections still empty after both extractors, fall back to text heuristic
        if not sections and result["text"]:
            result["warnings"].append("No heading-based sections found; using regex section detection")
            sections = _detect_sections_from_text(result["text"])

    else:
        raise HTTPException(status_code=422, detail=f"Unsupported file type: {file_type}")

    # ── Count words ───────────────────────────────────────
    word_count = len(result["text"].split())

    # ── Build metadata ─────────────────────────────────────
    metadata = ParseMetadata(
        file_id=file_id,
        original_filename=original_filename,
        file_type=file_type,
        page_count=result.get("page_count"),
        word_count=word_count,
        parsed_at=_now(),
        primary_extractor=result["extractor"],
        fallback_used=result["fallback_used"],
        fallback_extractor=result.get("fallback_extractor"),
        ocr_used=result.get("ocr_used", False),
        partial_parse=False,
        warnings=result.get("warnings", []),
    )

    # ── Build UnifiedDocumentJSON ─────────────────────────
    unified = UnifiedDocumentJSON(
        metadata=metadata,
        sections=sections,
        raw_text=result["text"],
        tables=tables,
    )

    # ── Save to parsed_docs/{user_id}/{file_id}.json ───────
    out_path = _parsed_path(user_id, file_id)
    out_path.write_text(
        unified.model_dump_json(indent=2),
        encoding="utf-8",
    )

    return ParseResponse(
        success=True,
        file_id=file_id,
        parsed_doc_path=str(out_path),
        metadata=metadata,
    )


async def get_parsed_document(user_id: str, file_id: str) -> UnifiedDocumentJSON:
    """Load and return an already-parsed document."""
    path = _parsed_path(user_id, file_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Parsed document not found. Parse it first.")
    return UnifiedDocumentJSON.model_validate_json(path.read_text(encoding="utf-8"))